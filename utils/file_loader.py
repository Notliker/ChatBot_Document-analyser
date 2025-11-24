import os
import tempfile
import streamlit as st
from pdf2image import convert_from_path
from pypdf import PdfReader
from docx import Document
from fpdf import FPDF

from langchain_core.documents import Document as langDocument
from utils.database import initialize_vector_db


DB_DOCS_LIMIT = 10


try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    print("RecursiveCharacterTextSplitter: импортирован успешно")
except ImportError:
    print("langchain_text_splitters не установлен или RecursiveCharacterTextSplitter не найден")


def load_pdf_to_db(pdf_file):
    """Загрузка PDF файла в векторную базу данных"""
    docs = []
    os.makedirs("docs", exist_ok=True)

    unique_sources = list(set(st.session_state.rag_sources))

    if pdf_file.name not in st.session_state.rag_sources:
        if len(unique_sources) < DB_DOCS_LIMIT:
            file_path = os.path.join("docs", pdf_file.name)
            with open(file_path, "wb") as f:
                f.write(pdf_file.read())

            try:
                loader = PDFProcessor()
                data = loader.process(file_path)

                for page_num, page_text in enumerate(data):
                    doc = langDocument(
                        page_content=page_text,
                        metadata={"source": pdf_file.name, "page": page_num + 1},
                    )
                    docs.append(doc)
                
                st.session_state.rag_sources.append(pdf_file.name)
                
                _split_and_load_docs(docs)
                
                st.toast(f"{pdf_file.name} успешно загружен")

            except Exception as e:
                st.toast(f"Ошибка загрузки PDF {pdf_file.name}: {e}")
                print(f"Ошибка загрузки PDF {pdf_file.name}: {e}")
        else:
            st.error(f"Достигнут лимит уникальных документов ({DB_DOCS_LIMIT}).")
            print(f"Лимит документов достигнут: {len(unique_sources)}/{DB_DOCS_LIMIT}")
    else:
        print(f"Документ {pdf_file.name} уже в базе")



def _split_and_load_docs(pages):
    '''Разделение текста на чанки (chunks)'''
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500, 
        chunk_overlap=150  
    )
    docs = text_splitter.split_documents(pages)

    if "vector_db" not in st.session_state or st.session_state.vector_db is None:
        st.session_state.vector_db = initialize_vector_db(pages)
    else:
        st.session_state.vector_db.add_documents(docs)



def handle_uploaded_files(files, session_rag_sources):
    session_rag_sources.extend(files)
    for file in files:
        if file.type == "application/pdf":
            load_pdf_to_db(file)
    return len(files)




class PDFProcessor:
    def __init__(self, ocr_model=None, captioner_model=None, min_text_length=20):
        from models.OCR import OCR, ImageCaptioner

        self.ocr = ocr_model or OCR()
        self.captioner = captioner_model or ImageCaptioner()
        self.temp_dir = tempfile.mkdtemp()
        self.min_text_length = min_text_length

    def is_text_file(self, file_path):
        """Check type file"""
        text_extensions = {'.txt', '.md', '.csv', '.log', '.text'}
        _, ext = os.path.splitext(file_path.lower())
        return ext in text_extensions

    def read_text_file(self, file_path):
        """Read text with differenf encodings"""
        try:
            encodings = ['utf-8', 'utf-8-sig', 'cp1251', 'windows-1251', 'latin-1', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    return text
                except (UnicodeDecodeError, LookupError):
                    continue
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return text
            
        except Exception as e:
            print(f"Ошибка чтения TXT файла: {e}")
            return ""

    def process_text_file(self, file_path):
        """Обработка TXT файла"""
        text = self.read_text_file(file_path)
        
        if not text:
            print("TXT файл пустой или не удалось прочитать")
            return []
        
        chunk_size = 2000
        pages = []
        
        if len(text) <= chunk_size:
            pages.append(text)
            print(f"  TXT файл обработан: 1 страница ({len(text)} символов)")
        else:
            lines = text.split('\n')
            current_chunk = []
            current_length = 0
            
            for line in lines:
                line_length = len(line) + 1  
                
                if current_length + line_length > chunk_size and current_chunk:
                    pages.append('\n'.join(current_chunk))
                    current_chunk = [line]
                    current_length = line_length
                else:
                    current_chunk.append(line)
                    current_length += line_length
            
            if current_chunk:
                pages.append('\n'.join(current_chunk))
            
            print(f"  TXT файл обработан: {len(pages)} страниц")
        
        return pages

    def convert_to_pdf(self, file_path):
        """Конвертация файлов в PDF """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".pdf":
            return file_path
        
        elif ext in ['.txt', '.md', '.csv', '.log']:
            return None
        
        elif ext == ".docx":
            pdf_path = file_path.replace(".docx", ".pdf")
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

            pdf = FPDF()
            pdf.add_page()
            pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
            pdf.set_font("DejaVu", size=12)
            for line in text.splitlines():
                pdf.cell(0, 10, text=line, new_x="LMARGIN", new_y="NEXT")
            pdf.output(pdf_path)
            return pdf_path
        
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def has_meaningful_text(self, text):
        """Check letter variety"""
        clean = text.strip()
        letters = sum(c.isalpha() for c in clean)
        return (
            len(clean) >= self.min_text_length and letters > self.min_text_length * 0.5
        )

    def image_extracter(self, page, page_num):
        """Extract images"""
        images_data = []

        try:
            for img_idx, image_file_object in enumerate(page.images):
                img_path = os.path.join(
                    self.temp_dir,
                    f"page_{page_num}_img_{img_idx}.{image_file_object.name.split('.')[-1]}",
                )

                with open(img_path, "wb") as f:
                    f.write(image_file_object.data)

                caption = self.captioner.describe(img_path)

                if caption:
                    images_data.append(
                        {"index": img_idx, "caption": caption, "path": img_path}
                    )

        except Exception as e:
            print(f"Error extracting images from page {page_num}: {e}")

        return images_data

    def format_images_text(self, images_data):
        """
        Convert image dictionary 2 string
        """
        if not images_data:
            return ""

        return "".join(
            [f"\n[IMAGE {img['index']}: {img['caption']}]" for img in images_data]
        )

    def process(self, file_path):
        """
        Full processing of file (PDF, DOCX, TXT)
        """
        if self.is_text_file(file_path):
            print(f"Обнаружен текстовый файл: {os.path.basename(file_path)}")
            return self.process_text_file(file_path)
        
        pdf_path = self.convert_to_pdf(file_path)
        
        if pdf_path is None:
            return self.process_text_file(file_path)
        
        processed_pages = []

        try:
            reader = PdfReader(pdf_path)
            try:
                pages_images = convert_from_path(pdf_path, dpi=200)
            except Exception as e:
                print(f"Warning: Could not convert to images: {e}")
                pages_images = []

            for page_num, page in enumerate(reader.pages):
                print(f"\nProcessing page {page_num + 1}.")
                text = page.extract_text()

                images_data = self.image_extracter(page, page_num)
                images_text = self.format_images_text(images_data)

                if self.has_meaningful_text(text):
                    final_text = text + images_text
                    processed_pages.append(final_text)
                    print(
                        f"  Used extracted text ({len(text)} chars) + {len(images_data)} images"
                    )

                else:
                    if page_num < len(pages_images):
                        img_path = os.path.join(
                            self.temp_dir, f"page_{page_num}_ocr.png"
                        )
                        pages_images[page_num].save(img_path, "PNG")

                        ocr_text = self.ocr.process(img_path)

                        if self.has_meaningful_text(ocr_text):
                            final_text = ocr_text + images_text
                            processed_pages.append(final_text)
                            print(
                                f"  Used OCR ({len(ocr_text)} chars) + {len(images_data)} images"
                            )
                        else:
                            caption = self.captioner.describe(img_path)
                            if caption:
                                final_text = (
                                    f"[FULL PAGE IMAGE: {caption}]" + images_text
                                )
                            else:
                                final_text = "[FULL PAGE IMAGE]" + images_text
                            processed_pages.append(final_text)
                            print(
                                f"  Used full page captioning + {len(images_data)} images"
                            )
                    else:
                        processed_pages.append(f"[NO TEXT EXTRACTED]{images_text}")
                        print(f"  No text, only {len(images_data)} embedded images")

        except Exception as e:
            print(f"Error processing PDF: {e}")
            raise
        
        return processed_pages

    def cleanup(self):
        """Очистка временных файлов"""
        try:
            for file in os.listdir(self.temp_dir):
                os.remove(os.path.join(self.temp_dir, file))
            os.rmdir(self.temp_dir)
            print("Cleanup completed")
        except Exception as e:
            print(f"Cleanup error: {e}")


